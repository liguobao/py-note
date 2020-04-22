import xlrd
import json
import os
from docx import Document


class ExcelInfo:
    # sheet表格名字，不填入默认取0
    sheet_name = None
    # 文件路径
    file_path = ""
    # 标题字段所在的列数组，会将对应列内容组装成标题
    title_fields = [1]
    # 答案选项开始列索引Id，包含当前列
    op_start_index = 2
    # 答案选项结束列索引Id，不包含当前列
    op_end_index = 7
    # 答案所在的列
    answer_index = 7


def convert_to_word(excel_info: ExcelInfo):
    file_name = excel_info.file_path if excel_info.file_path else "sample.xlsx"
    print("convert %s start!" % file_name)
    workbook_obj = xlrd.open_workbook(f'{file_name}')
    if excel_info.sheet_name:
        sheet_data = workbook_obj.sheet_by_name(excel_info.sheet_name)
    else:
        sheet_data = workbook_obj.sheet_by_index(0)
    data_num = 0
    document = Document()
    document.add_heading(file_name)
    for i in range(1, sheet_data.nrows):
        row_data = sheet_data.row_values(i)
        data_num = data_num + 1
        title_text = ""
        for title_index in excel_info.title_fields:
            title_text = title_text + row_data[title_index] + "  "
        title = f"{data_num}. {title_text}："
        options = []
        for op_id in range(excel_info.op_start_index, excel_info.op_end_index):
            op_value = row_data[op_id]
            if op_value is None or (type(op_value) != float and len(op_value) == 0):
                break
            options.append(op_value)
        paragraph = document.add_paragraph(title)
        for op_num in range(0, len(options)):
            op_text = "%s. %s" % (chr(65+op_num), options[op_num])
            document.add_paragraph(op_text)
        ans_text = row_data[excel_info.answer_index]
        ans = f"\n\n正确答案：【{ans_text}】\n\n"
        document.add_paragraph(ans)
        document.add_paragraph()
    document.save("%s.docx" % file_name)
    print("convert %s success!" % file_name)


excel_info = ExcelInfo()

convert_to_word(excel_info)
